from docx import Document

doc = Document()
doc.add_heading('Flute Roots - Technical Specifications & Architecture', 0)

doc.add_heading('1. Tech Stack & Architecture', level=2)
doc.add_paragraph('Frontend Framework: React 18 with TypeScript.', style='List Bullet')
doc.add_paragraph('Build Tool: Vite.', style='List Bullet')
doc.add_paragraph('Styling: Pure CSS (styles.css) utilizing a custom design system with CSS variables (nature-inspired palette: olives, greens, gold). It includes embedded SVG motifs (bamboo stalks, leaves) for visual aesthetics.', style='List Bullet')
doc.add_paragraph('Backend & Database: Supabase (PostgreSQL). It acts as your backend-as-a-service, handling both the database and user authentication.', style='List Bullet')
doc.add_paragraph('Routing: Custom state-based routing (AppRoute) inside App.tsx utilizing the HTML5 History API (window.history.pushState) rather than a heavy third-party router.', style='List Bullet')

doc.add_heading('2. Core Features & Routes', level=2)
doc.add_paragraph('The app features several primary views controlled by a state machine in App.tsx:')
doc.add_paragraph('Home (/): A landing page featuring an autoplaying hero audio snippet, an introductory YouTube video embed, and call-to-actions.', style='List Bullet')
doc.add_paragraph('Courses (/FluteRoots): A Learning Management System (LMS) interface where users can browse courses, see announcements, view the weekly class schedule, and enroll.', style='List Bullet')
doc.add_paragraph('Course Player (/coursePlayer): A dedicated interface for students to watch the video content of courses they are enrolled in.', style='List Bullet')
doc.add_paragraph('Biography (/biography): Detailed background on the artist, Digvijaysinh Chauhan.', style='List Bullet')
doc.add_paragraph('Organizers Corner (/organizersCorner): A portfolio/press-kit section tailored for event organizers containing a gallery, stage setup requirements, and availability calendar.', style='List Bullet')
doc.add_paragraph('Contact (/contact): Standard contact and booking information.', style='List Bullet')

doc.add_heading('3. Authentication & Admin Dashboard (/admin & /login)', level=2)
doc.add_paragraph('Authentication: Managed by Supabase Auth (supabase.auth.getSession()).', style='List Bullet')
doc.add_paragraph('Role-Based Access Control (RBAC): Admin rights are hardcoded to specific emails (digvijayflute@gmail.com and janhavikolekar280@gmail.com).', style='List Bullet')
doc.add_paragraph('Admin Dashboard: If an admin logs in, they get access to a powerful UI to dynamically manage courses, gallery images, dynamic text (announcements, bio paragraphs, awards), weekly schedules, events, and hero media.', style='List Bullet')

doc.add_heading('4. Database Schema (Supabase Tables Used)', level=2)
doc.add_paragraph('Based on the API calls, the app relies on 5 main tables:')
doc.add_paragraph('1. courses: Stores course metadata (title, description, price, video_url).', style='List Number')
doc.add_paragraph('2. enrollments: Maps users to courses they have access to (user_id, course_id).', style='List Number')
doc.add_paragraph('3. gallery: Stores image URLs and display ordering for the portfolio.', style='List Number')
doc.add_paragraph('4. events: Stores dates and types of events (performances, classes, blocked dates).', style='List Number')
doc.add_paragraph('5. settings: A flexible key-value store used to dynamically update site content (e.g., hero_image_url, bio_paragraphs_json, weekly_schedule_json) without needing a code redeploy.', style='List Number')

doc.add_heading('5. Notable Integrations & Utilities', level=2)
doc.add_paragraph('YouTube Parsing: Custom logic to automatically extract YouTube IDs from standard URLs to embed them natively via iframes.', style='List Bullet')
doc.add_paragraph('Image Processing: Integrates @imgly/background-removal for potential on-the-fly image manipulation.', style='List Bullet')
doc.add_paragraph('Local Caching: Extensively uses localStorage to cache settings, courses, and schedules so the site remains fast and doesn\'t flicker while waiting for Supabase to respond.', style='List Bullet')

doc.save('d:/A2gent_Work/Flute/PROJECT_SPECIFICATIONS.docx')
